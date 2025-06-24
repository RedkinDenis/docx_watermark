from docx import Document
from secret_messege import SecretMessage, Record

ZERO_WIDTH_NON_JOINER = "\u200C"  # Zero-width non-joiner (U+200C)
NORMAL_SPACE = "\u0020"      

def encode_zwbsp(doc_path, binary_str: str, output_path="secret.docx"):
    """
    Прячет бинарную строку в документ Word
    :param doc_path: Путь к исходному документу
    :param binary_str: Бинарная строка
    :param output_path: Путь для сохранения
    """
    doc = Document(doc_path)
    
    # Проверяем, что строка содержит только 0 и 1
    if not all(c in '01' for c in binary_str):
        raise ValueError("Binary string should contain only '0' and '1'")
    
    # Дополняем биты до чётного количества (на случай стороннего использования этой функции)
    if len(binary_str) % 2 != 0:
        binary_str += '0'
    
    bit_pairs = [binary_str[i:i+2] for i in range(0, len(binary_str), 2)]
    bit_index = 0

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if bit_index >= len(bit_pairs):
                break
            
            text = run.text
            new_text = []
            i = 0
            
            while i < len(text):
                if text[i] == NORMAL_SPACE and bit_index < len(bit_pairs):
                    bit_pair = bit_pairs[bit_index]
                    
                    if bit_pair == '00':
                        new_space = NORMAL_SPACE
                    elif bit_pair == '01':
                        new_space = NORMAL_SPACE + ZERO_WIDTH_NON_JOINER
                    elif bit_pair == '10':
                        new_space = ZERO_WIDTH_NON_JOINER + NORMAL_SPACE
                    elif bit_pair == '11':
                        new_space = ZERO_WIDTH_NON_JOINER + NORMAL_SPACE + ZERO_WIDTH_NON_JOINER
                    
                    new_text.append(new_space)
                    bit_index += 1
                else:
                    new_text.append(text[i])
                i += 1
            
            run.text = ''.join(new_text)
    
    doc.save(output_path)
    print(f"Бинарное сообщение внедрено в {output_path}")


def decode_zwbsp(doc_path) -> str:
    """
    Извлекает бинарную строку из документа Word
    :param doc_path: Путь к документу 
    :return: Бинарная строка 
    """
    doc = Document(doc_path)
    binary_secret = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text
        i = 0
        while i < len(text):
            if text[i] == NORMAL_SPACE:
                # Проверяем символы вокруг пробела
                prev_char = text[i-1] if i > 0 else ''
                next_char = text[i+1] if i < len(text)-1 else ''
                
                if prev_char == ZERO_WIDTH_NON_JOINER and next_char == ZERO_WIDTH_NON_JOINER:
                    binary_secret.append('11')
                elif prev_char == ZERO_WIDTH_NON_JOINER:
                    binary_secret.append('10')
                elif next_char == ZERO_WIDTH_NON_JOINER:
                    binary_secret.append('01')
                else:
                    binary_secret.append('00')
                i += 1
            i += 1
    
    return ''.join(binary_secret)

if __name__ == "__main__":

    import random
    random_bits = ''.join(random.choices('01', k=15))
    records = [
            Record(1, random_bits),
            Record(2, random_bits)
        ]
    msg = SecretMessage(records)
    
    print("Внедряемое сообщение")
    print(msg)
    binary_msg = msg.to_binary_str()
    
    encode_zwbsp("test.docx", binary_msg, "secret.docx")
    extracted_binary = decode_zwbsp("secret.docx")
    
    extr_msg = msg.from_binary_str(extracted_binary)
    
    print("Извлечённое сообщение")
    print(extr_msg)

    print(f"Внедрённое == извлечвенное: {msg == extr_msg}")
