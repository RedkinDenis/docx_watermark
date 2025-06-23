from docx import Document

Zero_Width_Non_Joiner='\u200C'
POPDirectional='\u202C'
Left_To_Right_Override='\u202D'
Left_To_Right_Mark='\u200E'
Right_To_Left_Override='\u202E'
Narrow_No_Break_Space='\u202F'
Left_To_rightembedding='\u202A'
Right_To_leftembedding='\u202B'
Mongolian_vowelseparator='\u180E'
Right_To_LeftMark='\u200F'
Zero_Width_Joiner='\u200D'
Zero_Width_Space='\u200B'
Zero_Width_Non_Break='\uFEFF'
HairSpace='\u200A'
Six_Per_EmSpace='\u2006'
FigureSpace='\u2007'
PunctuationSpace='\u2008'
ThinSpace='\u2009'
EnQuad='\u2000'
Three_Per_EmSpace='\u2004'
Four_Per_EmSpace='\u2005'
NormalSpace='\u0020'

unicode_chars = [
    ("Zero_Width_Non_Joiner", '\u200C', '200C'),
    ("POPDirectional", '\u202C', '202C'),
    ("Left_To_Right_Override", '\u202D', '202D'),
    ("Left_To_Right_Mark", '\u200E', '200E'),
    ("Right_To_Left_Override", '\u202E', '202E'),
    ("Narrow_No_Break_Space", '\u202F', '202F'),
    ("Left_To_rightembedding", '\u202A', '202A'),
    ("Right_To_leftembedding", '\u202B', '202B'),
    ("Mongolian_vowelseparator", '\u180E', '180E'),
    ("Right_To_LeftMark", '\u200F', '200F'),
    ("Zero_Width_Joiner", '\u200D', '200D'),
    ("Zero_Width_Space", '\u200B', '200B'),
    ("Zero_Width_Non_Break", '\uFEFF', 'FEFF'),
    ("HairSpace", '\u200A', '200A'),
    ("Six_Per_EmSpace", '\u2006', '2006'),
    ("FigureSpace", '\u2007', '2007'),
    ("PunctuationSpace", '\u2008', '2008'),
    ("ThinSpace", '\u2009', '2009'),
    ("EnQuad", '\u2000', '2000'),
    ("Three_Per_EmSpace", '\u2004', '2004'),
    ("Four_Per_EmSpace", '\u2005', '2005'),
    ("NormalSpace", '\u0020', '0020')
]

string =  "\n".join("[" + char[2] + "]" + char[0] + ":(" + char[1] + ")" for char in unicode_chars) 
# print(string)

doc = Document()

doc.add_paragraph()
doc.paragraphs[0].text = string

doc.save("unicode_chars.docx")





